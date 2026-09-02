from __future__ import annotations

import csv
import io
import json
import os
from typing import Protocol

from dataguard.processing.models import DocumentType, ExtractedDocument
from dataguard.processing.validation import UnsafeDocumentError


class Extractor(Protocol):
    def extract(self, filename: str, content: bytes) -> ExtractedDocument: ...


def _clean(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.replace("\x00", "").splitlines()).strip()


class TextExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        try:
            text = content.decode("utf-8-sig", errors="strict")
        except UnicodeDecodeError as exc:
            raise UnsafeDocumentError("text document is not valid UTF-8") from exc
        return ExtractedDocument(filename, DocumentType.TXT, _clean(text))


class JSONExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        try:
            value = json.loads(content.decode("utf-8-sig"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise UnsafeDocumentError("JSON document is malformed or not valid UTF-8") from exc
        return ExtractedDocument(
            filename, DocumentType.JSON, _clean(json.dumps(value, ensure_ascii=False, indent=2))
        )


class CSVExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        try:
            text = content.decode("utf-8-sig", errors="strict")
        except UnicodeDecodeError as exc:
            raise UnsafeDocumentError("CSV document is not valid UTF-8") from exc
        rows = csv.reader(io.StringIO(text))
        normalized = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
        return ExtractedDocument(filename, DocumentType.CSV, _clean(normalized))


class PDFExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError

        try:
            reader = PdfReader(io.BytesIO(content), strict=True)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except (PdfReadError, ValueError, OSError) as exc:
            raise UnsafeDocumentError("PDF document could not be parsed") from exc
        return ExtractedDocument(filename, DocumentType.PDF, _clean(text), len(reader.pages))


class DOCXExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = Document(io.BytesIO(content))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        except (PackageNotFoundError, ValueError, KeyError, OSError) as exc:
            raise UnsafeDocumentError("DOCX document could not be parsed") from exc
        return ExtractedDocument(filename, DocumentType.DOCX, _clean("\n".join(parts)))


class XLSXExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        from openpyxl import load_workbook
        from openpyxl.utils.exceptions import InvalidFileException

        try:
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        except (InvalidFileException, ValueError, KeyError, OSError) as exc:
            raise UnsafeDocumentError("XLSX document could not be parsed") from exc
        try:
            parts: list[str] = []
            for sheet in workbook.worksheets:
                parts.append(f"[SHEET] {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    parts.append(" | ".join("" if value is None else str(value) for value in row))
            return ExtractedDocument(filename, DocumentType.XLSX, _clean("\n".join(parts)))
        except (ValueError, KeyError, OSError) as exc:
            raise UnsafeDocumentError("XLSX document could not be parsed") from exc
        finally:
            workbook.close()


class ImageExtractor:
    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        from PIL import Image, UnidentifiedImageError

        try:
            image = Image.open(io.BytesIO(content))
            image.verify()
        except (UnidentifiedImageError, OSError, ValueError) as exc:
            raise UnsafeDocumentError("image document could not be parsed") from exc
        if os.getenv("DATAGUARD_OCR_ENABLED", "false").lower() != "true":
            return ExtractedDocument(
                filename,
                DocumentType.IMAGE,
                "",
                warnings=(
                    "OCR is disabled; set DATAGUARD_OCR_ENABLED=true to extract image text.",
                ),
            )
        from dataguard.processing.ocr import OCRExtractor

        return OCRExtractor().extract(filename, content)
